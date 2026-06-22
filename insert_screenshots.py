#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Insert captured screenshots into the Sara manual Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MANUAL = Path(r"D:\Work\Free lance\Sara\New App - Sara\Sara_Abo_El_Ela_Architecture_Office_System_Manual.docx")
SCREENSHOTS = Path(r"D:\Work\Free lance\Sara\New App - Sara\screenshots")
OUT = MANUAL

MAP = {
    "شاشة تسجيل الدخول": "login.png",
    "قائمة التنقل السفلية": "dashboard.png",
    "نموذج إضافة عميل": "modal_add_client.png",
    "نموذج إضافة مشروع": "modal_add_project.png",
    "نموذج مصروف مشروع": "modal_project_expense.png",
    "نموذج مشتريات": "modal_procurement.png",
    "نموذج إضافة عهدة": "modal_office_custody.png",
}

PLACEHOLDER_RE = re.compile(r"\[صورة:\s*(.+?)\]")


def clear_paragraph(p):
    for run in p.runs:
        run.text = ""


def set_run_font(run, size=11, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def main():
    doc = Document(str(MANUAL))
    replaced = 0
    for p in doc.paragraphs:
        text = p.text
        m = PLACEHOLDER_RE.search(text)
        if not m:
            continue
        caption = m.group(1).strip()
        filename = MAP.get(caption)
        if not filename:
            print(f"No screenshot mapped for: {caption}")
            continue
        img_path = SCREENSHOTS / filename
        if not img_path.exists():
            print(f"Missing screenshot file: {img_path}")
            continue
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.right_to_left = True
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.8))
        cap_run = p.add_run(f"\n{caption}")
        set_run_font(cap_run, size=10, color=(100, 100, 100))
        replaced += 1
        print(f"Inserted {filename} for '{caption}'")

    doc.save(str(OUT))
    print(f"\nDone. Replaced {replaced} placeholders.")
    print(f"Saved to: {OUT}")


if __name__ == "__main__":
    main()
