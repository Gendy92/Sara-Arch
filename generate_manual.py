#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate the Sara Abo El-Ela Architecture Office System user manual as a .docx file.
Target audience: non-technical Arabic-speaking users (~55 years old).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_PATH = r"D:\Work\Free lance\Sara\New App - Sara\Sara_Abo_El_Ela_Architecture_Office_System_Manual.docx"

# Helper to set run font for Arabic support
def set_run_font(run, size=14, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Complex script font (Arabic) in Word is set via the eastAsia/ascs settings,
    # but setting name on the run is usually enough for display.
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/drawingml/2006/main}cs", "Arial")


def add_rtl_paragraph(doc, text, size=14, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.right_to_left = True
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_rtl(doc, text, level=1):
    # Use Word heading styles, then fix direction
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size=(20 if level == 1 else (16 if level == 2 else 14)), bold=True, color=(180, 140, 70) if level == 1 else None)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet_rtl(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True
    p.paragraph_format.space_after = Pt(4)
    # Adjust left/right indent for nested bullets
    p.paragraph_format.right_indent = Inches(0.25 * indent_level)
    run = p.add_run(text)
    set_run_font(run, size=13)
    return p


def add_numbered_rtl(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13)
    return p


def add_note(doc, text):
    p = add_rtl_paragraph(doc, "💡 ملاحظة: " + text, size=12, color=(80, 80, 80), spacing_after=8)
    return p


def add_warning(doc, text):
    p = add_rtl_paragraph(doc, "⚠️ تنبيه: " + text, size=12, color=(200, 60, 40), spacing_after=8)
    return p


def add_step(doc, number, title, details):
    p = add_rtl_paragraph(doc, f"الخطوة {number}: {title}", size=13, bold=True, spacing_after=2)
    p2 = add_rtl_paragraph(doc, details, size=13, spacing_after=8)
    return p, p2


def add_image_placeholder(doc, caption):
    p = add_rtl_paragraph(doc, f"[صورة: {caption}]", size=11, color=(120, 120, 120), alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)
    return p


def main():
    doc = Document()

    # Document-wide default font for new paragraphs
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(14)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Cover page
    add_rtl_paragraph(doc, "", size=20, spacing_after=60)
    add_rtl_paragraph(doc, "Sara Abo El-Ela Architecture Office System Manual", size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    add_rtl_paragraph(doc, "دليل استخدام نظام مكتب سارة أبو العلا للمحاسبة والإدارة المالية", size=24, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=24)
    add_rtl_paragraph(doc, "إصدار 1.0", size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    add_rtl_paragraph(doc, "تم إعداده للمستخدمين غير التقنيين", size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=60)
    doc.add_page_break()

    # Table of contents (manual)
    add_heading_rtl(doc, "فهرس الدليل", level=1)
    toc = [
        "1. قبل أن تبدأ",
        "2. التسجيل وقائمة التنقل",
        "3. لوحة التحكم",
        "4. العملاء والمشاريع",
        "5. معاملات المشاريع",
        "6. الموردين والمشتريات",
        "7. حساب المكتب",
        "8. العهد النقدية",
        "9. الموظفين",
        "10. المهام",
        "11. البيانات الأساسية",
        "12. الإعدادات والأمان",
        "13. النسخ الاحتياطي والصيانة",
        "14. حلول المشاكل الشائعة",
        "15. معجم المصطلحات",
    ]
    for item in toc:
        add_bullet_rtl(doc, item)
    doc.add_page_break()

    # 1. Before you start
    add_heading_rtl(doc, "1. قبل أن تبدأ", level=1)
    add_rtl_paragraph(doc, "هذا الدليل مكتوب بخطوات بسيطة. اقرأ الفصل كاملاً قبل أن تبدأ بالضغط على الأزرار، ثم ارجع إليه كلما احتجت مساعدة.", spacing_after=8)
    add_heading_rtl(doc, "1.1 ما هو نظام سارة؟", level=2)
    add_rtl_paragraph(doc, "هو برنامج على الإنترنت يساعد مكتب سارة أبو العلا في تسجيل:", spacing_after=4)
    add_bullet_rtl(doc, "إيرادات المشاريع والمكتب")
    add_bullet_rtl(doc, "مصروفات المشاريع والمكتب")
    add_bullet_rtl(doc, "الموردين والمشتريات")
    add_bullet_rtl(doc, "الموظفين والرواتب والحضور")
    add_bullet_rtl(doc, "العهد النقدية")
    add_bullet_rtl(doc, "المهام لكل مشروع")
    add_heading_rtl(doc, "1.2 مصطلحات مهمة يجب معرفتها", level=2)
    add_bullet_rtl(doc, "العميل: الشخص أو الشركة صاحبة المشروع.")
    add_bullet_rtl(doc, "المشروع: عملية تصميم أو إشراف معينة لعميل.")
    add_bullet_rtl(doc, "المورد: الشخص أو الشركة التي ندفع لها مقابل مواد أو خدمات.")
    add_bullet_rtl(doc, "معاملة: أي سجل مالي (إيداع، مصروف، سحب، ...).")
    add_bullet_rtl(doc, "العهدة النقدية: مبلغ نقدي نسلمه لموظف لإنفاقه نيابة عن المكتب.")
    add_bullet_rtl(doc, "نسبة الإشراف: نسبة من قيمة أعمال المشروع تُحسب كإيراد للمكتب.")
    add_bullet_rtl(doc, "حساب المكتب: يجمع النقدية والبنك والإيرادات والمصروفات العامة.")
    add_note(doc, "كل الأرقام في النظام بالجنية المصري (ج.م) ما لم يُذكر خلاف ذلك.")
    doc.add_page_break()

    # 2. Login and navigation
    add_heading_rtl(doc, "2. التسجيل وقائمة التنقل", level=1)
    add_heading_rtl(doc, "2.1 كيف تفتح النظام", level=2)
    add_step(doc, 1, "افتح المتصفح", "استخدم Google Chrome أو Microsoft Edge.")
    add_step(doc, 2, "اكتب عنوان الموقع", "الرابط يعطيه لك المسؤول (مثلاً: https://gendy92.github.io/Sara-Arch).")
    add_step(doc, 3, "سجل الدخول", "اكتب البريد الإلكتروني وكلمة المرور، ثم اضغط 'تسجيل الدخول'.")
    add_image_placeholder(doc, "شاشة تسجيل الدخول")
    add_heading_rtl(doc, "2.2 قائمة التنقل", level=2)
    add_rtl_paragraph(doc, "بعد تسجيل الدخول تظهر قائمة في أسفل الشاشة (أو جانبها على الكمبيوتر) تحتوي على الأيقونات:", spacing_after=4)
    add_bullet_rtl(doc, "🏠 لوحة التحكم: ملخص سريع لكل شيء.")
    add_bullet_rtl(doc, "👥 العملاء: قائمة العملاء والمشاريع.")
    add_bullet_rtl(doc, "💰 معاملات المشاريع: إيداعات ومصروفات المشاريع.")
    add_bullet_rtl(doc, "🏢 المكتب: حساب المكتب والعهد.")
    add_bullet_rtl(doc, "🚚 الموردين: الموردين والمشتريات.")
    add_bullet_rtl(doc, "🧑‍💼 الموظفين: الموظفين والرواتب.")
    add_bullet_rtl(doc, "📋 المهام: متابعة مهام المشاريع.")
    add_bullet_rtl(doc, "⚙️ الإعدادات: بيانات المكتب والمستخدمين.")
    add_image_placeholder(doc, "قائمة التنقل السفلية")
    doc.add_page_break()

    # 3. Dashboard
    add_heading_rtl(doc, "3. لوحة التحكم", level=1)
    add_rtl_paragraph(doc, "هذه هي الواجهة الرئيسية بعد الدخول. تعرض أهم الأرقام والرسوم البيانية.", spacing_after=8)
    add_heading_rtl(doc, "3.1 بطاقات المؤشرات (KPIs)", level=2)
    add_rtl_paragraph(doc, "في أعلى الصفحة توجد بطاقات ملونة. كل بطاقة تعرض رقماً واحداً:", spacing_after=4)
    add_bullet_rtl(doc, "إجمالي الإيرادات")
    add_bullet_rtl(doc, "إجمالي المصروفات")
    add_bullet_rtl(doc, "صافي الربح (الإيرادات ناقص المصروفات)")
    add_bullet_rtl(doc, "رصيد المكتب (النقدي + البنك)")
    add_bullet_rtl(doc, "عدد المشاريع والعملاء والموظفين")
    add_note(doc, "يمكنك الضغط على بعض البطاقات للانتقال إلى الشاشة التفصيلية.")
    add_heading_rtl(doc, "3.2 الرسوم البيانية", level=2)
    add_bullet_rtl(doc, "مشاريعيًا: إيرادات ومصروفات المشاريع شهريًا.")
    add_bullet_rtl(doc, "مكتبيًا: إيرادات ومصروفات المكتب شهريًا.")
    add_bullet_rtl(doc, "توزيع مصروفات المكتب: حسب التصنيف.")
    add_heading_rtl(doc, "3.3 قوائم التنبيهات", level=2)
    add_bullet_rtl(doc, "أرصدة الموردين المستحقة: مبالغ يجب دفعها للموردين.")
    add_bullet_rtl(doc, "أرصدة العملاء النشطين: مبالغ مستحقة من العملاء.")
    add_bullet_rtl(doc, "عهد نقدية غير مقفلة: عهد لها رصيد متبقي.")
    doc.add_page_break()

    # 4. Clients and projects
    add_heading_rtl(doc, "4. العملاء والمشاريع", level=1)
    add_heading_rtl(doc, "4.1 إضافة عميل جديد", level=2)
    add_step(doc, 1, "انتقل إلى شاشة العملاء", "اضغط 'العملاء' من القائمة السفلية.")
    add_step(doc, 2, "اضغط 'إضافة عميل'", "زر أزرق في أعلى الصفحة.")
    add_step(doc, 3, "املأ البيانات", "الاسم (مطلوب)، الهاتف، البريد، العنوان، الملاحظات.")
    add_step(doc, 4, "اضغط 'حفظ الكل'", "[صورة: نموذج إضافة عميل]")
    add_heading_rtl(doc, "4.2 إضافة مشروع لعميل", level=2)
    add_step(doc, 1, "افتح تفاصيل العميل", "اضغط على اسم العميل في القائمة.")
    add_step(doc, 2, "اضغط 'إضافة مشروع'", "في صفحة تفاصيل العميل.")
    add_step(doc, 3, "املأ بيانات المشروع", "الاسم (مطلوب)، القيمة الإجمالية، نسبة الإشراف، نسبة التصميم، تاريخ البدء والانتهاء، العنوان، الملاحظات.")
    add_step(doc, 4, "اضغط 'حفظ الكل'", "[صورة: نموذج إضافة مشروع]")
    add_note(doc, "نسبة الإشراف تُستخدم لاحقًا لحساب إيراد المكتب تلقائيًا.")
    add_heading_rtl(doc, "4.3 تفاصيل المشروع", level=2)
    add_rtl_paragraph(doc, "عند فتح مشروع ترى:", spacing_after=4)
    add_bullet_rtl(doc, "بطاقات: القيمة، الإيداعات، المصروفات، الإشراف، الرصيد.")
    add_bullet_rtl(doc, "زر 'كشف حساب': لتحميل ملف Excel لتفاصيل المشروع.")
    add_bullet_rtl(doc, "زر 'مهام': لإدارة مهام المشروع.")
    add_bullet_rtl(doc, "زر 'مرتجع': لتسجيل مرتجع للعميل.")
    add_bullet_rtl(doc, "جدول المعاملات: كل الإيداعات والمصروفات الخاصة بالمشروع.")
    add_bullet_rtl(doc, "جدول المهام: مهام المشروع وحالتها.")
    doc.add_page_break()

    # 5. Project transactions
    add_heading_rtl(doc, "5. معاملات المشاريع", level=1)
    add_rtl_paragraph(doc, "هذه المعاملات تؤثر مباشرة على رصيد المشروع. توجد في شاشة 'معاملات المشاريع'.", spacing_after=8)
    add_heading_rtl(doc, "5.1 عربون مشروع (إيداع من عميل)", level=2)
    add_step(doc, 1, "اضغط '💰 عربون مشروع'", "من أعلى شاشة المعاملات.")
    add_step(doc, 2, "اختر العميل ثم المشروع", "لا يمكن اختيار المشروع قبل العميل.")
    add_step(doc, 3, "اكتب المبلغ", "المبلغ الذي دفعه العميل.")
    add_step(doc, 4, "اختر طريقة الدفع", "نقدي / بنكي / تحويل.")
    add_step(doc, 5, "اكتب التاريخ والبيان", "التاريخ مطلوب، البيان اختياري.")
    add_step(doc, 6, "اضغط 'حفظ الكل'", "")
    add_heading_rtl(doc, "5.2 مصروف مشروع", level=2)
    add_step(doc, 1, "اضغط '🔨 مصروف مشروع'", "")
    add_step(doc, 2, "اختر العميل ثم المشروع", "")
    add_step(doc, 3, "اختر المورد", "إذا كان المورد هو 'مكتب سارة أبو العلا' سيظهر تنبيه قبل الحفظ.")
    add_step(doc, 4, "اختر القسم والبند", "اختياري، يساعد في التصنيف.")
    add_step(doc, 5, "اكتب المبلغ والمدفوع", "إذا كان المدفوع أقل من المبلغ، تُسجل المعاملة كآجلة.")
    add_step(doc, 6, "اختر طريقة الدفع والتاريخ والبيان", "")
    add_step(doc, 7, "اضغط 'حفظ الكل'", "[صورة: نموذج مصروف مشروع]")
    add_warning(doc, "لا تستخدم 'مكتب سارة أبو العلا' كمورد لمصاريف الإشراف أو التصميم أو إيجار المعدات. هذه تُسجل كإيرادات مكتب من شاشة 'حساب المكتب'.")
    add_heading_rtl(doc, "5.3 مرتجع عميل", level=2)
    add_step(doc, 1, "من تفاصيل المشروع اضغط '⬅️ مرتجع'", "")
    add_step(doc, 2, "اختر المشروع والعميل", "")
    add_step(doc, 3, "اكتب المبلغ والتاريخ والبيان", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "5.4 تسديد متأخرات مورد", level=2)
    add_step(doc, 1, "من تفاصيل المورد اضغط 'تسديد' أو من شاشة المعاملات", "")
    add_step(doc, 2, "اختر المورد والمشروع", "")
    add_step(doc, 3, "اكتب المبلغ والتاريخ والبيان", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "5.5 الإشراف على مشروع", level=2)
    add_rtl_paragraph(doc, "إذا أردت تسجيل إشراف يدوي:", spacing_after=4)
    add_step(doc, 1, "اضغط 'إشراف مشروع'", "")
    add_step(doc, 2, "اختر المشروع", "")
    add_step(doc, 3, "اكتب المبلغ والتاريخ والبيان", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_note(doc, "عادةً يتم حساب الإشراف تلقائيًا من نسبة الإشراف المسجلة في المشروع، لكن يمكنك استخدام هذا الزر للإشرافات الخارجة عن القاعدة.")
    add_heading_rtl(doc, "5.6 فلترة المعاملات", level=2)
    add_bullet_rtl(doc, "تبويب 'الكل': جميع المعاملات.")
    add_bullet_rtl(doc, "تبويب 'المصروفات': المصروفات فقط.")
    add_bullet_rtl(doc, "أزرار 'الكل / الإيداعات / المصروفات': فلترة سريعة داخل التبويب.")
    doc.add_page_break()

    # 6. Vendors and procurements
    add_heading_rtl(doc, "6. الموردين والمشتريات", level=1)
    add_heading_rtl(doc, "6.1 إضافة مورد", level=2)
    add_step(doc, 1, "انتقل إلى 'الموردين'", "")
    add_step(doc, 2, "اضغط 'إضافة مورد'", "")
    add_step(doc, 3, "املأ الاسم (مطلوب)، الهاتف، البريد، العنوان، نوع المورد (خدمة/بضاعة)، القطاع، الملاحظات", "")
    add_step(doc, 4, "اضغط 'حفظ الكل'", "")
    add_heading_rtl(doc, "6.2 إضافة مشتريات", level=2)
    add_step(doc, 1, "من شاشة الموردين أو تفاصيل المورد اضغط 'إضافة مشتريات'", "")
    add_step(doc, 2, "اختر المورد والعميل والمشروع", "")
    add_step(doc, 3, "اكتب البند / الصنف (مطلوب)", "")
    add_step(doc, 4, "اكتب الكمية وسعر الوحدة", "الإجمالي = الكمية × سعر الوحدة")
    add_step(doc, 5, "اختر التصنيف (مواد / عمالة / معدات / أخرى)", "")
    add_step(doc, 6, "اكتب التاريخ والملاحظات", "")
    add_step(doc, 7, "اضغط 'حفظ الكل'", "[صورة: نموذج مشتريات]")
    add_heading_rtl(doc, "6.3 كشف حساب المورد", level=2)
    add_rtl_paragraph(doc, "من تفاصيل المورد اضغط 'كشف حساب' لتحميل Excel يوضح جميع المعاملات والمشتريات والرصيد.", spacing_after=8)
    add_heading_rtl(doc, "6.4 تنبيه مكتب سارة كمورد", level=2)
    add_rtl_paragraph(doc, "إذا اخترت 'مكتب سارة أبو العلا' كمورد في أي نموذج، سيظهر التنبيه التالي قبل الحفظ:", spacing_after=4)
    add_rtl_paragraph(doc, "⚠️ تنبيه: هذا المورد هو المكتب الداخلي. اختيار سارة كمورد قد يؤدي إلى ازدواجية في التكلفة، حيث أن مصاريف المكتب مسجلة بالفعل في النظام. تأكد أن هذا المصروف ليس تكلفة إشراف أو تصميم أو إيجار معدات — هذه تُسجل كإيرادات مكتب وليس كمصروف مشروع. هل تريد المتابعة؟", size=12, color=(180, 60, 40), spacing_after=8)
    add_rtl_paragraph(doc, "اضغط 'متابعة' إذا كنت متأكدًا، أو 'إلغاء' لاختيار مورد آخر.", spacing_after=8)
    doc.add_page_break()

    # 7. Office account
    add_heading_rtl(doc, "7. حساب المكتب", level=1)
    add_rtl_paragraph(doc, "شاشة 'حساب المكتب' تجمع كل ما يخص النقدية والبنك والمصروفات العامة والإيرادات.", spacing_after=8)
    add_heading_rtl(doc, "7.1 بطاقات المكتب", level=2)
    add_bullet_rtl(doc, "رصيد نقدي: المبلغ الموجود نقدًا.")
    add_bullet_rtl(doc, "رصيد بنكي: المبلغ الموجود في البنك.")
    add_bullet_rtl(doc, "رصيد المكتب الإجمالي: النقدي + البنك + إيرادات الإشراف.")
    add_bullet_rtl(doc, "العهد النقدية: مجموع العهد المفتوحة.")
    add_heading_rtl(doc, "7.2 مصروف مكتبي", level=2)
    add_step(doc, 1, "اضغط '🏢 مصروف مكتبي'", "")
    add_step(doc, 2, "اختر الموظف والتصنيف", "")
    add_step(doc, 3, "اختر المورد (اختياري)", "")
    add_step(doc, 4, "اكتب المبلغ واختر الحساب (نقدي/بنكي)", "")
    add_step(doc, 5, "اكتب التاريخ والبيان", "")
    add_step(doc, 6, "اضغط 'حفظ الكل'", "")
    add_heading_rtl(doc, "7.3 إيراد مكتبي", level=2)
    add_step(doc, 1, "اضغط '📈 إيراد مكتبي'", "")
    add_step(doc, 2, "اكتب المبلغ واختر الحساب", "")
    add_step(doc, 3, "اختر التصنيف (مثلاً: تصميم، إيجار معدات، أخرى)", "")
    add_step(doc, 4, "اكتب التاريخ والبيان", "مثال: 'تصميم فيلا العميل أحمد' أو 'إيجار معدات مشروع الزراعة'")
    add_step(doc, 5, "اضغط 'حفظ الكل'", "")
    add_note(doc, "استخدم 'إيراد مكتبي' لتسجيل إيرادات التصميم وإيجار المعدات وأي إيراد عام للمكتب.")
    add_heading_rtl(doc, "7.4 توريد / سحب صاحب المكتب", level=2)
    add_bullet_rtl(doc, "توريد صاحب المكتب: عندما يضع صاحب المكتب مالًا شخصيًا في حساب المكتب.")
    add_bullet_rtl(doc, "سحب صاحب المكتب: عندما يأخذ صاحب المكتب مالًا من حساب المكتب لاستخدامه الشخصي.")
    add_heading_rtl(doc, "7.5 تحميل Excel لحساب المكتب", level=2)
    add_step(doc, 1, "اضغط '📥 تحميل Excel'", "في أعلى شاشة المكتب.")
    add_step(doc, 2, "سيُحفظ الملف على جهازك", "يحتوي على تفاصيل المعاملات.")
    doc.add_page_break()

    # 8. Custody
    add_heading_rtl(doc, "8. العهد النقدية", level=1)
    add_heading_rtl(doc, "8.1 ما هي العهدة النقدية؟", level=2)
    add_rtl_paragraph(doc, "العهدة هي مبلغ نقدي تسلمه لموظف لإنفاقه نيابة عن المكتب. النظام يراقب المبلغ الأصلي والمصروفات والمرتجع والباقي.", spacing_after=8)
    add_heading_rtl(doc, "8.2 إضافة عهدة نقدية", level=2)
    add_step(doc, 1, "من شاشة المكتب اضغط '💼 عهد نقدية'", "")
    add_step(doc, 2, "اختر الموظف", "")
    add_step(doc, 3, "اكتب المبلغ", "مبلغ العهدة الأصلي.")
    add_step(doc, 4, "اختر الحساب (نقدي/بنكي)", "")
    add_step(doc, 5, "اكتب التاريخ والملاحظات", "")
    add_step(doc, 6, "اضغط 'حفظ الكل'", "[صورة: نموذج إضافة عهدة]")
    add_heading_rtl(doc, "8.3 إضافة مصروف عهدة", level=2)
    add_step(doc, 1, "اضغط '🔨 مصروف عهدة'", "من شاشة المكتب.")
    add_step(doc, 2, "اختر نوع المصروف أولاً", "مكتب أو مشروع. حتى تختار النوع، لن يتم تفعيل قائمة العهد.")
    add_step(doc, 3, "اختر العهدة", "تظهر العهد المفتوحة فقط مع الرصيد المتاح.")
    add_step(doc, 4, "املأ بيانات المصروف", "حسب النوع:")
    add_bullet_rtl(doc, "مكتب: الموظف (مقفل)، التصنيف، المورد، المبلغ، الحساب، التاريخ، البيان.", indent_level=1)
    add_bullet_rtl(doc, "مشروع: الموظف (مقفل)، العميل، المشروع، المورد، القسم، البند، المبلغ، المدفوع، طريقة الدفع، التاريخ، البيان.", indent_level=1)
    add_step(doc, 5, "اضغط 'حفظ الكل'", "")
    add_warning(doc, "لا يمكن تجاوز الرصيد المتاح للعهدة. إذا حاولت ذلك، سيظهر لك تنبيه.")
    add_heading_rtl(doc, "8.4 سداد باقي العهدة", level=2)
    add_step(doc, 1, "اضغط على 'مصروفات' بجانب العهدة في جدول العهد", "")
    add_step(doc, 2, "اضغط '💵 سداد باقي'", "")
    add_step(doc, 3, "اكتب المبلغ المرتجع والتاريخ والبيان", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "8.5 حالات العهدة", level=2)
    add_bullet_rtl(doc, "مفتوحة (active): لم يُنفق منها شيء بعد.")
    add_bullet_rtl(doc, "جزئية (partial): تم إنفاق جزء والباقي أكبر من صفر.")
    add_bullet_rtl(doc, "مقفلة (settled): تم إنفاقها بالكامل أو تم تسجيل المرتجع.")
    doc.add_page_break()

    # 9. Employees
    add_heading_rtl(doc, "9. الموظفين", level=1)
    add_heading_rtl(doc, "9.1 إضافة موظف", level=2)
    add_step(doc, 1, "انتقل إلى 'الموظفين'", "")
    add_step(doc, 2, "اضغط 'إضافة موظفين'", "")
    add_step(doc, 3, "املأ الاسم (مطلوب)، الهاتف، البريد، العنوان، تاريخ التعيين، الملاحظات", "")
    add_step(doc, 4, "اضغط 'حفظ الكل'", "")
    add_heading_rtl(doc, "9.2 معاملات الموظف", level=2)
    add_rtl_paragraph(doc, "من تفاصيل الموظف يمكن إضافة:", spacing_after=4)
    add_bullet_rtl(doc, "سلفة (advance): مبلغ أخذه الموظف مقدمًا من الراتب.")
    add_bullet_rtl(doc, "جزاء (penalty): خصم من الراتب.")
    add_bullet_rtl(doc, "مكافأة (bonus): مبلغ إضافي.")
    add_bullet_rtl(doc, "أخرى (other): أي معاملة أخرى.")
    add_step(doc, 1, "افتح تفاصيل الموظف", "")
    add_step(doc, 2, "اضغط 'إضافة معاملة'", "")
    add_step(doc, 3, "اختر النوع والمبلغ والتاريخ والبيان", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "9.3 رفع ملف البصمة", level=2)
    add_step(doc, 1, "في شاشة الموظفين، مرر للأسفل إلى 'رفع ملف البصمة'", "")
    add_step(doc, 2, "اختر الشهر والسنة", "")
    add_step(doc, 3, "اختر ملف Excel/CSV من جهازك", "")
    add_step(doc, 4, "تأكد من المعاينة ثم اضغط 'معالجة البصمة'", "")
    add_step(doc, 5, "انتقل إلى 'الرواتب الشهرية' لمراجعتها", "")
    add_heading_rtl(doc, "9.4 الرواتب الشهرية", level=2)
    add_step(doc, 1, "اختر الشهر والسنة", "")
    add_step(doc, 2, "راجع الأيام والخصومات والإضافات", "")
    add_step(doc, 3, "اضغط 'حفظ الرواتب'", "")
    add_note(doc, "تأكد من رفع البصمة أولاً إذا كنت تستخدم نظام الحضور والانصراف.")
    add_heading_rtl(doc, "9.5 سجل الرواتب", level=2)
    add_rtl_paragraph(doc, "يمكنك تعديل الراتب الأساسي للموظف من 'تعديل موظف'، وستُسجل التغييرات تلقائيًا في سجل الرواتب.", spacing_after=8)
    doc.add_page_break()

    # 10. Tasks
    add_heading_rtl(doc, "10. المهام", level=1)
    add_heading_rtl(doc, "10.1 إضافة مهمة لمشروع", level=2)
    add_step(doc, 1, "افتح تفاصيل المشروع", "")
    add_step(doc, 2, "اضغط '📋 مهام'", "")
    add_step(doc, 3, "اضغط 'إضافة مهمة'", "")
    add_step(doc, 4, "اكتب اسم المهمة والمسؤول وتاريخ البدء والاستحقاق", "")
    add_step(doc, 5, "اختر الحالة (معلق/قيد التنفيذ/منتهي) والأولوية", "")
    add_step(doc, 6, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "10.2 شاشة المهام", level=2)
    add_rtl_paragraph(doc, "من القائمة السفلية اضغط '📋 المهام' لرؤية جميع المهام في كل المشاريع. يمكنك الفلترة حسب الحالة.", spacing_after=8)
    add_heading_rtl(doc, "10.3 تعديل أو حذف مهمة", level=2)
    add_bullet_rtl(doc, "اضغط 'تعديل' لتغيير البيانات.")
    add_bullet_rtl(doc, "اضغط 'حذف' لإزالتها. سيُطلب تأكيد قبل الحذف.")
    doc.add_page_break()

    # 11. Master data
    add_heading_rtl(doc, "11. البيانات الأساسية", level=1)
    add_rtl_paragraph(doc, "البيانات الأساسية هي القوائم التي تستخدمها في النماذج (التصنيفات، الأصناف، أقسام المشاريع، بنود الأعمال).", spacing_after=8)
    add_heading_rtl(doc, "11.1 الوصول إلى البيانات الأساسية", level=2)
    add_step(doc, 1, "انتقل إلى 'الإعدادات'", "")
    add_step(doc, 2, "اضغط 'البيانات الأساسية'", "")
    add_heading_rtl(doc, "11.2 التصنيفات", level=2)
    add_bullet_rtl(doc, "تُستخدم في مصروفات المكتب والإيرادات.")
    add_bullet_rtl(doc, "مثال: مصاريف إدارية، مواصلات، كهرباء، إيجار.")
    add_heading_rtl(doc, "11.3 الأصناف / المواد", level=2)
    add_bullet_rtl(doc, "قائمة المواد التي قد تشتريها للمشاريع.")
    add_bullet_rtl(doc, "تُستخدم عند إضافة المشتريات.")
    add_heading_rtl(doc, "11.4 أقسام المشاريع", level=2)
    add_bullet_rtl(doc, "تصنيفات أعمال المشروع مثل: إنشاءات، تشطيبات، كهرباء، سباكة.")
    add_heading_rtl(doc, "11.5 بنود الأعمال", level=2)
    add_bullet_rtl(doc, "بنود تفصيلية ضمن كل قسم، مثل: 'أعمال خرسانة' أو 'دهانات'.")
    add_bullet_rtl(doc, "تُستخدم في مصروفات المشاريع لمعرفة نوع العمل المُنفق عليه.")
    doc.add_page_break()

    # 12. Settings and security
    add_heading_rtl(doc, "12. الإعدادات والأمان", level=1)
    add_heading_rtl(doc, "12.1 إعدادات المكتب", level=2)
    add_step(doc, 1, "انتقل إلى 'الإعدادات'", "")
    add_step(doc, 2, "عدل البيانات التالية حسب الحاجة", "")
    add_bullet_rtl(doc, "اسم الشركة / المكتب", indent_level=1)
    add_bullet_rtl(doc, "العنوان والهاتف والرقم الضريبي", indent_level=1)
    add_bullet_rtl(doc, "نسبة الإشراف الافتراضية (%)", indent_level=1)
    add_bullet_rtl(doc, "تسمية العملة", indent_level=1)
    add_step(doc, 3, "اضغط '💾 حفظ الإعدادات'", "")
    add_heading_rtl(doc, "12.2 المستخدمين والصلاحيات", level=2)
    add_rtl_paragraph(doc, "متاح فقط للمدير (Admin).", spacing_after=4)
    add_step(doc, 1, "من الإعدادات اضغط 'فتح المستخدمين'", "")
    add_step(doc, 2, "اضغط 'إضافة مستخدمين'", "")
    add_step(doc, 3, "املأ اسم المستخدم، الاسم الكامل، كلمة المرور، الدور", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "12.3 الصلاحيات", level=2)
    add_rtl_paragraph(doc, "يمكن للمدير تحديد كل مستخدم ما إذا كان يستطيع عرض أو إضافة أو تعديل أو حذف في كل شاشة.", spacing_after=8)
    add_step(doc, 1, "من شاشة المستخدمين اضغط 'صلاحيات المستخدمين'", "")
    add_step(doc, 2, "اختر المستخدم", "")
    add_step(doc, 3, "حدد الصلاحيات لكل شاشة", "")
    add_step(doc, 4, "اضغط 'حفظ'", "")
    add_heading_rtl(doc, "12.4 سجل العمليات", level=2)
    add_rtl_paragraph(doc, "يسجل النظام كل إضافة وتعديل وحذف. من 'الإعدادات > سجل العمليات' يمكنك مراجعة من قام بكل عملية ومتى.", spacing_after=8)
    doc.add_page_break()

    # 13. Backup and maintenance
    add_heading_rtl(doc, "13. النسخ الاحتياطي والصيانة", level=1)
    add_heading_rtl(doc, "13.1 تحميل نسخة احتياطية", level=2)
    add_step(doc, 1, "انتقل إلى 'النسخ الاحتياطي'", "")
    add_step(doc, 2, "اضغط '📥 تحميل النسخة الاحتياطية'", "")
    add_step(doc, 3, "سيُحفظ ملف ZIP على جهازك", "يحتوي على جميع البيانات.")
    add_note(doc, "احفظ الملف في مكان آمن (مثل فلاشة أو Google Drive) وكرر العملية بشكل دوري.")
    add_heading_rtl(doc, "13.2 مسح الكاش وإعادة التحميل", level=2)
    add_rtl_paragraph(doc, "إذا لم تظهر التحديثات الجديدة، أو ظهرت بيانات قديمة:", spacing_after=4)
    add_step(doc, 1, "انتقل إلى 'النسخ الاحتياطي'", "")
    add_step(doc, 2, "اضغط '🧹 مسح الكاش وإعادة التحميل'", "")
    add_step(doc, 3, "انتظر حتى يعيد المتصفح فتح التطبيق", "")
    doc.add_page_break()

    # 14. Troubleshooting
    add_heading_rtl(doc, "14. حلول المشاكل الشائعة", level=1)
    add_heading_rtl(doc, "14.1 لا تظهر بيانات جديدة", level=2)
    add_bullet_rtl(doc, "اضغط F5 لتحديث الصفحة.")
    add_bullet_rtl(doc, "إذا استمرت المشكلة، اذهب إلى 'النسخ الاحتياطي' واضغط 'مسح الكاش وإعادة التحميل'.")
    add_heading_rtl(doc, "14.2 لا يمكن اختيار المشروع", level=2)
    add_bullet_rtl(doc, "تأكد أنك اخترت العميل أولاً.")
    add_bullet_rtl(doc, "تأكد أن العميل له مشروع مسجل.")
    add_heading_rtl(doc, "14.3 مصروف العهدة لا يُحفظ", level=2)
    add_bullet_rtl(doc, "تأكد أن العهدة مفتوحة ولها رصيد متاح.")
    add_bullet_rtl(doc, "تأكد أن إجمالي المصروفات لا يتجاوز الرصيد المتاح.")
    add_heading_rtl(doc, "14.4 نسيت كلمة المرور", level=2)
    add_bullet_rtl(doc, "اتصل بالمدير أو المسؤول الفني لإعادة تعيين كلمة المرور.")
    add_heading_rtl(doc, "14.5 ظهر خطأ 'closeDropdown is not defined'", level=2)
    add_bullet_rtl(doc, "هذا خطأ قديم في القوائم المنسدلة. قم بمسح الكاش وإعادة التحميل.")
    add_heading_rtl(doc, "14.6 لا أجد زر 'حفظ'", level=2)
    add_bullet_rtl(doc, "في النماذج الطويلة، قد يكون الزر أسفل الشاشة. مرر للأسفل.")
    add_bullet_rtl(doc, "تأكد أنك ملأت جميع الحقول المطلوبة (التي تحمل علامة *).")
    doc.add_page_break()

    # 15. Glossary
    add_heading_rtl(doc, "15. معجم المصطلحات", level=1)
    terms = [
        ("العميل", "الشخص أو الشركة صاحبة المشروع."),
        ("المشروع", "عملية تصميم أو إشراف محددة."),
        ("المورد", "من ندفع له مقابل مواد أو خدمات."),
        ("المعاملة", "سجل مالي واحد (إيداع، مصروف، ...)."),
        ("العهدة", "مبلغ نقدي تسلمه لموظف لإنفاقه."),
        ("الإشراف", "إيراد المكتب من إشرافه على أعمال المشروع."),
        ("الرصيد", "المبلغ المتبقي بعد طرح المصروفات من الإيداعات."),
        ("الحساب النقدي", "الأموال الفعلية في الخزنة."),
        ("الحساب البنكي", "الأموال في البنك."),
        ("المسوّى", "ما تم إنفاقه أو تسديده من العهدة."),
        ("العهد المقفلة", "العهدة التي تم إنفاقها أو تسديد باقيها بالكامل."),
    ]
    for term, definition in terms:
        p = add_rtl_paragraph(doc, f"{term}: {definition}", size=13, spacing_after=6)
        # Bold the term
        p.runs[0].font.bold = True
        # Re-apply font name after bold
        p.runs[0].font.name = "Arial"

    # Footer note
    doc.add_paragraph()
    add_rtl_paragraph(doc, "--- نهاية الدليل ---", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    add_rtl_paragraph(doc, "للمساعدة، اتصل بالمسؤول الفني أو مراجع هذا الدليل.", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)

    doc.save(OUTPUT_PATH)
    print(f"Manual saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
